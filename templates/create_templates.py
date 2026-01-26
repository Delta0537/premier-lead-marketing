#!/usr/bin/env python3
"""
Create standardized PLM templates: Proposal and Competitive Analysis
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    for edge in ['top', 'left', 'bottom', 'right']:
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def create_proposal_template():
    """Create the main proposal/quote template"""
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # ========== HEADER SECTION ==========
    # Company Header
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.columns[0].width = Inches(4)
    header_table.columns[1].width = Inches(2.5)

    # Logo cell
    logo_cell = header_table.rows[0].cells[0]
    logo_para = logo_cell.paragraphs[0]
    logo_run = logo_para.add_run("PLM")
    logo_run.bold = True
    logo_run.font.size = Pt(36)
    logo_run.font.color.rgb = RGBColor(59, 130, 246)  # Blue

    tagline = logo_cell.add_paragraph("PREMIER LEAD MARKETING")
    tagline.runs[0].font.size = Pt(14)
    tagline.runs[0].font.color.rgb = RGBColor(20, 184, 166)  # Teal
    tagline.runs[0].bold = True

    subtitle = logo_cell.add_paragraph("Strategic Digital Dominance")
    subtitle.runs[0].font.size = Pt(10)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    subtitle.runs[0].italic = True

    # Contact cell
    contact_cell = header_table.rows[0].cells[1]
    contact_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    contact_para = contact_cell.paragraphs[0]
    contact_para.add_run("andrew@premierleadmarketing.com\n").font.size = Pt(9)
    contact_para.add_run("premierleadmarketing.com\n").font.size = Pt(9)
    contact_para.add_run("Premier Lead Marketing, LLC").font.size = Pt(9)

    doc.add_paragraph()

    # Blue divider line
    divider = doc.add_paragraph()
    divider_run = divider.add_run("_" * 85)
    divider_run.font.color.rgb = RGBColor(59, 130, 246)

    # ========== PROPOSAL TITLE ==========
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("MARKETING SERVICES PROPOSAL")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(30, 30, 30)

    doc.add_paragraph()

    # ========== CLIENT INFORMATION TABLE ==========
    info_heading = doc.add_paragraph()
    info_run = info_heading.add_run("CLIENT INFORMATION")
    info_run.bold = True
    info_run.font.size = Pt(12)
    info_run.font.color.rgb = RGBColor(59, 130, 246)

    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'

    info_data = [
        ("Prepared For:", "[Client Name]"),
        ("Company:", "[Company Name]"),
        ("Date:", "[Date]"),
        ("Valid Until:", "[30 days from date]")
    ]

    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        info_table.rows[i].cells[1].text = value
        set_cell_shading(info_table.rows[i].cells[0], 'F3F4F6')

    doc.add_paragraph()

    # ========== EXECUTIVE SUMMARY ==========
    exec_heading = doc.add_paragraph()
    exec_run = exec_heading.add_run("EXECUTIVE SUMMARY")
    exec_run.bold = True
    exec_run.font.size = Pt(12)
    exec_run.font.color.rgb = RGBColor(59, 130, 246)

    exec_para = doc.add_paragraph()
    exec_para.add_run("[Brief overview of the client's challenges, goals, and how PLM will deliver results. Keep this to 2-3 sentences that speak directly to their pain points and desired outcomes.]")
    exec_para.paragraph_format.left_indent = Inches(0.25)

    doc.add_paragraph()

    # ========== PROPOSED SERVICES SECTION ==========
    services_heading = doc.add_paragraph()
    services_run = services_heading.add_run("PROPOSED SERVICES")
    services_run.bold = True
    services_run.font.size = Pt(14)
    services_run.font.color.rgb = RGBColor(59, 130, 246)

    instruction = doc.add_paragraph()
    instruction_run = instruction.add_run("[Select applicable services below - delete rows not needed]")
    instruction_run.italic = True
    instruction_run.font.size = Pt(9)
    instruction_run.font.color.rgb = RGBColor(150, 150, 150)

    # Service items with checkboxes - each as a separate table for modularity
    services = [
        {
            "name": "Lead Generation Services",
            "description": "Precision-targeted lead generation campaigns using Apollo & Clay integration",
            "features": [
                "Multi-channel outreach campaigns",
                "Quality-verified leads delivered to your CRM",
                "Custom targeting based on your ideal customer profile",
                "Weekly lead delivery reports"
            ],
            "price": "$____/month"
        },
        {
            "name": "GoHighLevel Automation Setup",
            "description": "Complete end-to-end automation platform setup and management",
            "features": [
                "Full GHL system setup & configuration",
                "Custom funnels & landing pages",
                "Email & SMS automation sequences",
                "Calendar & booking system integration"
            ],
            "price": "$____ setup + $____/month"
        },
        {
            "name": "Marketing Systems & Campaigns",
            "description": "Multi-channel marketing infrastructure for lead nurturing",
            "features": [
                "Automated follow-up sequences",
                "Multi-channel campaign management",
                "Performance tracking & analytics",
                "A/B testing & optimization"
            ],
            "price": "$____/month"
        },
        {
            "name": "Website Development",
            "description": "Professional, high-converting website design and development",
            "features": [
                "Custom responsive website design",
                "Mobile-optimized layouts",
                "Marketing integration ready",
                "SEO foundation setup"
            ],
            "price": "$____ one-time"
        },
        {
            "name": "Client Portal Development",
            "description": "Custom client portal for streamlined operations",
            "features": [
                "Branded client dashboard",
                "Document management",
                "Communication hub",
                "Progress tracking"
            ],
            "price": "$____ one-time"
        },
        {
            "name": "Business Coaching & Strategy",
            "description": "Expert guidance for scaling your business",
            "features": [
                "Monthly strategy sessions",
                "Competitive market analysis",
                "Growth planning",
                "1-on-1 coaching calls"
            ],
            "price": "$____/month"
        },
        {
            "name": "Unified Marketing Portal Access",
            "description": "Single dashboard for all marketing metrics and performance",
            "features": [
                "Real-time analytics dashboard",
                "Campaign performance tracking",
                "Lead pipeline visualization",
                "ROI reporting & insights"
            ],
            "price": "Included with automation package"
        }
    ]

    for service in services:
        doc.add_paragraph()

        # Service box
        service_table = doc.add_table(rows=1, cols=1)
        service_table.style = 'Table Grid'
        cell = service_table.rows[0].cells[0]

        # Service name
        name_para = cell.paragraphs[0]
        checkbox_run = name_para.add_run("[ ] ")
        checkbox_run.font.size = Pt(12)
        name_run = name_para.add_run(service["name"])
        name_run.bold = True
        name_run.font.size = Pt(12)
        name_run.font.color.rgb = RGBColor(30, 30, 30)

        # Description
        desc_para = cell.add_paragraph()
        desc_run = desc_para.add_run(service["description"])
        desc_run.font.size = Pt(10)
        desc_run.font.color.rgb = RGBColor(80, 80, 80)

        # Features
        for feature in service["features"]:
            feat_para = cell.add_paragraph()
            feat_para.paragraph_format.left_indent = Inches(0.25)
            feat_run = feat_para.add_run(f"• {feature}")
            feat_run.font.size = Pt(10)

        # Price
        price_para = cell.add_paragraph()
        price_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        price_label = price_para.add_run("Investment: ")
        price_label.font.size = Pt(10)
        price_value = price_para.add_run(service["price"])
        price_value.bold = True
        price_value.font.size = Pt(11)
        price_value.font.color.rgb = RGBColor(59, 130, 246)

    doc.add_paragraph()
    doc.add_paragraph()

    # ========== INVESTMENT SUMMARY ==========
    invest_heading = doc.add_paragraph()
    invest_run = invest_heading.add_run("INVESTMENT SUMMARY")
    invest_run.bold = True
    invest_run.font.size = Pt(14)
    invest_run.font.color.rgb = RGBColor(59, 130, 246)

    # Summary table
    summary_table = doc.add_table(rows=5, cols=2)
    summary_table.style = 'Table Grid'
    summary_table.columns[0].width = Inches(4.5)
    summary_table.columns[1].width = Inches(2)

    summary_data = [
        ("One-Time Setup Fees:", "$____"),
        ("Monthly Retainer:", "$____"),
        ("", ""),
        ("TOTAL MONTHLY INVESTMENT:", "$____"),
        ("Contract Term:", "[Month-to-month / 3 months / 6 months]")
    ]

    for i, (item, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = item
        summary_table.rows[i].cells[1].text = value
        summary_table.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if i == 3:  # Total row
            summary_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            summary_table.rows[i].cells[1].paragraphs[0].runs[0].bold = True
            set_cell_shading(summary_table.rows[i].cells[0], 'EBF5FF')
            set_cell_shading(summary_table.rows[i].cells[1], 'EBF5FF')

    doc.add_paragraph()

    # ========== TIMELINE ==========
    time_heading = doc.add_paragraph()
    time_run = time_heading.add_run("PROJECT TIMELINE")
    time_run.bold = True
    time_run.font.size = Pt(12)
    time_run.font.color.rgb = RGBColor(59, 130, 246)

    timeline_table = doc.add_table(rows=5, cols=2)
    timeline_table.style = 'Table Grid'

    timeline_data = [
        ("Phase", "Timeline"),
        ("Discovery & Audit", "Week 1"),
        ("Strategy & Planning", "Week 2"),
        ("Implementation", "Weeks 3-4"),
        ("Optimize & Scale", "Ongoing")
    ]

    for i, (phase, time) in enumerate(timeline_data):
        timeline_table.rows[i].cells[0].text = phase
        timeline_table.rows[i].cells[1].text = time
        if i == 0:
            timeline_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            timeline_table.rows[i].cells[1].paragraphs[0].runs[0].bold = True
            set_cell_shading(timeline_table.rows[i].cells[0], '3B82F6')
            set_cell_shading(timeline_table.rows[i].cells[1], '3B82F6')
            timeline_table.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            timeline_table.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()

    # ========== TERMS & CONDITIONS ==========
    terms_heading = doc.add_paragraph()
    terms_run = terms_heading.add_run("TERMS & CONDITIONS")
    terms_run.bold = True
    terms_run.font.size = Pt(12)
    terms_run.font.color.rgb = RGBColor(59, 130, 246)

    terms = [
        "Payment is due upon signing. Monthly retainer billed on the 1st of each month.",
        "30-day notice required for service cancellation after initial term.",
        "All deliverables remain property of PLM until final payment is received.",
        "Client agrees to provide necessary access, content, and feedback in a timely manner.",
        "Results are not guaranteed but our team is committed to achieving measurable outcomes."
    ]

    for term in terms:
        term_para = doc.add_paragraph()
        term_para.add_run(f"• {term}").font.size = Pt(10)

    doc.add_paragraph()

    # ========== SIGNATURE SECTION ==========
    sig_heading = doc.add_paragraph()
    sig_run = sig_heading.add_run("AGREEMENT")
    sig_run.bold = True
    sig_run.font.size = Pt(12)
    sig_run.font.color.rgb = RGBColor(59, 130, 246)

    agree_para = doc.add_paragraph()
    agree_para.add_run("By signing below, both parties agree to the terms outlined in this proposal.")

    doc.add_paragraph()

    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.columns[0].width = Inches(3.25)
    sig_table.columns[1].width = Inches(3.25)

    # Client signature
    sig_table.rows[0].cells[0].text = "CLIENT:"
    sig_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    sig_table.rows[1].cells[0].text = "\n\n________________________________"
    sig_table.rows[2].cells[0].text = "Signature"
    sig_table.rows[3].cells[0].text = "Date: _______________"

    # PLM signature
    sig_table.rows[0].cells[1].text = "PREMIER LEAD MARKETING:"
    sig_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    sig_table.rows[1].cells[1].text = "\n\n________________________________"
    sig_table.rows[2].cells[1].text = "Andrew Knight, COO"
    sig_table.rows[3].cells[1].text = "Date: _______________"

    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Premier Lead Marketing, LLC | Strategic Digital Dominance")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(150, 150, 150)

    # Save document
    doc.save('/home/user/premier-lead-marketing/templates/PLM_Proposal_Template.docx')
    print("Created: PLM_Proposal_Template.docx")


def create_competitive_analysis_template():
    """Create competitive analysis template"""
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # ========== HEADER ==========
    header_table = doc.add_table(rows=1, cols=2)
    header_table.columns[0].width = Inches(4)
    header_table.columns[1].width = Inches(2.5)

    logo_cell = header_table.rows[0].cells[0]
    logo_para = logo_cell.paragraphs[0]
    logo_run = logo_para.add_run("PLM")
    logo_run.bold = True
    logo_run.font.size = Pt(36)
    logo_run.font.color.rgb = RGBColor(59, 130, 246)

    tagline = logo_cell.add_paragraph("PREMIER LEAD MARKETING")
    tagline.runs[0].font.size = Pt(14)
    tagline.runs[0].font.color.rgb = RGBColor(20, 184, 166)
    tagline.runs[0].bold = True

    contact_cell = header_table.rows[0].cells[1]
    contact_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    contact_para = contact_cell.paragraphs[0]
    contact_para.add_run("andrew@premierleadmarketing.com\n").font.size = Pt(9)
    contact_para.add_run("premierleadmarketing.com").font.size = Pt(9)

    doc.add_paragraph()

    # Divider
    divider = doc.add_paragraph()
    divider_run = divider.add_run("_" * 85)
    divider_run.font.color.rgb = RGBColor(59, 130, 246)

    # ========== TITLE ==========
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("COMPETITIVE MARKET ANALYSIS")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(30, 30, 30)

    doc.add_paragraph()

    # ========== CLIENT INFO ==========
    info_heading = doc.add_paragraph()
    info_run = info_heading.add_run("ANALYSIS OVERVIEW")
    info_run.bold = True
    info_run.font.size = Pt(12)
    info_run.font.color.rgb = RGBColor(59, 130, 246)

    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'

    info_data = [
        ("Client:", "[Client Name]"),
        ("Industry:", "[Industry]"),
        ("Market:", "[Geographic Market / Niche]"),
        ("Analysis Date:", "[Date]"),
        ("Prepared By:", "Premier Lead Marketing")
    ]

    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        info_table.rows[i].cells[1].text = value
        set_cell_shading(info_table.rows[i].cells[0], 'F3F4F6')

    doc.add_paragraph()

    # ========== EXECUTIVE SUMMARY ==========
    exec_heading = doc.add_paragraph()
    exec_run = exec_heading.add_run("EXECUTIVE SUMMARY")
    exec_run.bold = True
    exec_run.font.size = Pt(12)
    exec_run.font.color.rgb = RGBColor(59, 130, 246)

    exec_table = doc.add_table(rows=1, cols=1)
    exec_table.style = 'Table Grid'
    exec_cell = exec_table.rows[0].cells[0]
    exec_cell.text = "[2-3 paragraph summary of key findings, competitive landscape, and strategic recommendations]"
    set_cell_shading(exec_cell, 'F9FAFB')

    doc.add_paragraph()

    # ========== COMPETITOR OVERVIEW ==========
    comp_heading = doc.add_paragraph()
    comp_run = comp_heading.add_run("COMPETITOR OVERVIEW")
    comp_run.bold = True
    comp_run.font.size = Pt(14)
    comp_run.font.color.rgb = RGBColor(59, 130, 246)

    # Competitor comparison table
    comp_table = doc.add_table(rows=6, cols=5)
    comp_table.style = 'Table Grid'

    headers = ["Competitor", "Services", "Pricing", "Strengths", "Weaknesses"]
    for i, header in enumerate(headers):
        comp_table.rows[0].cells[i].text = header
        comp_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        comp_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(comp_table.rows[0].cells[i], '3B82F6')

    for row in range(1, 6):
        for col in range(5):
            if col == 0:
                comp_table.rows[row].cells[col].text = f"[Competitor {row}]"
            else:
                comp_table.rows[row].cells[col].text = "[...]"

    doc.add_paragraph()

    # ========== INDIVIDUAL COMPETITOR ANALYSIS ==========
    # Template for 3 competitors
    for comp_num in range(1, 4):
        doc.add_paragraph()

        comp_detail = doc.add_paragraph()
        comp_detail_run = comp_detail.add_run(f"COMPETITOR {comp_num}: [Name]")
        comp_detail_run.bold = True
        comp_detail_run.font.size = Pt(12)
        comp_detail_run.font.color.rgb = RGBColor(59, 130, 246)

        detail_table = doc.add_table(rows=7, cols=2)
        detail_table.style = 'Table Grid'
        detail_table.columns[0].width = Inches(2)
        detail_table.columns[1].width = Inches(4.5)

        details = [
            ("Website:", "[URL]"),
            ("Services Offered:", "[List key services]"),
            ("Pricing Model:", "[Pricing structure]"),
            ("Target Market:", "[Who they serve]"),
            ("Marketing Channels:", "[How they acquire customers]"),
            ("Key Differentiators:", "[What makes them unique]"),
            ("Vulnerabilities:", "[Where you can compete]")
        ]

        for i, (label, value) in enumerate(details):
            detail_table.rows[i].cells[0].text = label
            detail_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            detail_table.rows[i].cells[1].text = value
            set_cell_shading(detail_table.rows[i].cells[0], 'F3F4F6')

    doc.add_paragraph()

    # ========== MARKET POSITIONING ==========
    pos_heading = doc.add_paragraph()
    pos_run = pos_heading.add_run("MARKET POSITIONING MATRIX")
    pos_run.bold = True
    pos_run.font.size = Pt(14)
    pos_run.font.color.rgb = RGBColor(59, 130, 246)

    pos_table = doc.add_table(rows=5, cols=4)
    pos_table.style = 'Table Grid'

    pos_headers = ["Factor", "Client", "Avg Competitor", "Opportunity"]
    for i, header in enumerate(pos_headers):
        pos_table.rows[0].cells[i].text = header
        pos_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        pos_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(pos_table.rows[0].cells[i], '3B82F6')

    factors = ["Online Presence", "Service Quality", "Price Point", "Customer Experience"]
    for i, factor in enumerate(factors):
        pos_table.rows[i+1].cells[0].text = factor
        pos_table.rows[i+1].cells[1].text = "[1-10]"
        pos_table.rows[i+1].cells[2].text = "[1-10]"
        pos_table.rows[i+1].cells[3].text = "[High/Med/Low]"

    doc.add_paragraph()

    # ========== SWOT ANALYSIS ==========
    swot_heading = doc.add_paragraph()
    swot_run = swot_heading.add_run("CLIENT SWOT ANALYSIS")
    swot_run.bold = True
    swot_run.font.size = Pt(14)
    swot_run.font.color.rgb = RGBColor(59, 130, 246)

    swot_table = doc.add_table(rows=2, cols=2)
    swot_table.style = 'Table Grid'

    swot_data = [
        ("STRENGTHS\n\n• [Strength 1]\n• [Strength 2]\n• [Strength 3]",
         "WEAKNESSES\n\n• [Weakness 1]\n• [Weakness 2]\n• [Weakness 3]"),
        ("OPPORTUNITIES\n\n• [Opportunity 1]\n• [Opportunity 2]\n• [Opportunity 3]",
         "THREATS\n\n• [Threat 1]\n• [Threat 2]\n• [Threat 3]")
    ]

    colors = [('DCFCE7', 'FEE2E2'), ('DBEAFE', 'FEF3C7')]  # Green/Red, Blue/Yellow

    for i, (left, right) in enumerate(swot_data):
        swot_table.rows[i].cells[0].text = left
        swot_table.rows[i].cells[1].text = right
        set_cell_shading(swot_table.rows[i].cells[0], colors[i][0])
        set_cell_shading(swot_table.rows[i].cells[1], colors[i][1])

    doc.add_paragraph()

    # ========== STRATEGIC RECOMMENDATIONS ==========
    rec_heading = doc.add_paragraph()
    rec_run = rec_heading.add_run("STRATEGIC RECOMMENDATIONS")
    rec_run.bold = True
    rec_run.font.size = Pt(14)
    rec_run.font.color.rgb = RGBColor(59, 130, 246)

    rec_table = doc.add_table(rows=6, cols=3)
    rec_table.style = 'Table Grid'

    rec_headers = ["Priority", "Recommendation", "Expected Impact"]
    for i, header in enumerate(rec_headers):
        rec_table.rows[0].cells[i].text = header
        rec_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        rec_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(rec_table.rows[0].cells[i], '3B82F6')

    priorities = ["HIGH", "HIGH", "MEDIUM", "MEDIUM", "LOW"]
    for i, priority in enumerate(priorities):
        rec_table.rows[i+1].cells[0].text = priority
        rec_table.rows[i+1].cells[1].text = f"[Recommendation {i+1}]"
        rec_table.rows[i+1].cells[2].text = "[Impact description]"

        if priority == "HIGH":
            set_cell_shading(rec_table.rows[i+1].cells[0], 'FEE2E2')
        elif priority == "MEDIUM":
            set_cell_shading(rec_table.rows[i+1].cells[0], 'FEF3C7')
        else:
            set_cell_shading(rec_table.rows[i+1].cells[0], 'DCFCE7')

    doc.add_paragraph()

    # ========== NEXT STEPS ==========
    next_heading = doc.add_paragraph()
    next_run = next_heading.add_run("RECOMMENDED NEXT STEPS")
    next_run.bold = True
    next_run.font.size = Pt(12)
    next_run.font.color.rgb = RGBColor(59, 130, 246)

    steps = [
        "Schedule strategy session to discuss findings",
        "Prioritize quick-win opportunities",
        "Develop differentiation strategy",
        "Create implementation roadmap",
        "Begin execution of Phase 1 recommendations"
    ]

    for i, step in enumerate(steps):
        step_para = doc.add_paragraph()
        step_para.add_run(f"{i+1}. {step}")

    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Premier Lead Marketing, LLC | Strategic Digital Dominance")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(150, 150, 150)

    confidential = doc.add_paragraph()
    confidential.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_run = confidential.add_run("CONFIDENTIAL - For intended recipient only")
    conf_run.font.size = Pt(8)
    conf_run.font.color.rgb = RGBColor(200, 200, 200)
    conf_run.italic = True

    # Save document
    doc.save('/home/user/premier-lead-marketing/templates/PLM_Competitive_Analysis_Template.docx')
    print("Created: PLM_Competitive_Analysis_Template.docx")


if __name__ == "__main__":
    print("Creating PLM Templates...")
    print("-" * 40)
    create_proposal_template()
    create_competitive_analysis_template()
    print("-" * 40)
    print("All templates created successfully!")
